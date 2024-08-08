import os
from typing import Self
import PyPDF2
import numpy as np
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from typing import List

from preprocessor import ImageSkewCorrector, Binarization, NoiseRemoval
from layout_detector import LayoutDetector

class PdfToImageConvertor:
    """
    Converts a pdf file to an png image.

    Attributes:
        input_file: a pdf file. 
    """
    def __init__(self, input_file):
        """
        Initializes import of a pdf file.

        Args:
        input_file: a pdf file. 
        """
        self.input_file = input_file

    def pdf_to_images(self, output_folder: str) -> List[str]:
        """
        Converts a pdf file to an image PNG.

        Args: 
            output_folder: a folder to save an image.

        Returns:
            image_paths: a list of an image path.
        """
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        images = convert_from_path(self.input_file)
        image_paths = []
        for i, image in enumerate(images):
            image_path = f'{output_folder}/page_{i + 1}.png'
            image.save(image_path)
            image_paths.append(image_path)
        return image_paths

class ImagePreProcessor:
    """
    Implements preprocessor.py.

    Attributes:
        image_paths: a list of an image path. 
    """
    # initializing import an image
    def __init__(self, image_path: str):
        """
        Initializes an image_path.

        Args:
        image_paths: a list of an image path.  
        """
        self.image_path = image_path

    def preprocess_image(self) -> np.ndarray:
        """
        Preprocess and saves an image.

        Returns:
            final_image_array: an image (NumPy array format) after preprocessing.
        """ 
        image = Image.open(self.image_path)

        # Converting image to numpy array for processing
        image_np = np.array(image)

        # Correcting Skew
        skew_corrector = ImageSkewCorrector(image_np)
        corrected_image = skew_corrector.correct_skew()

        # Converting to Gray and Binarize
        binarizer = Binarization(corrected_image)
        gray_image = binarizer.gray_conversion()
        binarized_image = binarizer.binarized_conversion(gray_image)

        # Removing Noise
        noise_removal = NoiseRemoval(binarized_image)
        final_image_array = noise_removal.gaussian_blurring()
        return final_image_array

class Layout:
    """
    Implements layout_detector.py

    Attributes:
        image_array: an image (NumPy array format) after preprocessing. 
    """
    def __init__(self, final_image_array: np.ndarray):
        """
        Initializes an image.

        Args:
            image_array: an image (NumPy array format) after preprocessing. 
        """
        self.image_array = final_image_array
    
    def layout_detection(self) -> dict:
        """
        Detects layout in a preprocessed image.

        Returns: 
            extracted_text: a dictionarry of text blocks and their names.
        """
        layout_detector = LayoutDetector(self.image_array)
        extracted_text = layout_detector.process_image_layout()
        return extracted_text

class Doc():
    """
    Saves extracted text to the .docx file.

    Attributes: 
        exracted_text: a dictionarry of text blocks and their names.

    """
    def __init__(self, extracted_text: dict):
        """
        Initializes an image.

        Args:
            extracted_text: a dictionarry of text blocks and their names.
        """
        self.input_text = extracted_text

    def creating_doc(self):
        """
        Saves extracted text to the .docx file.
        """
        doc = Document()
        for block_type, blocks in self.input_text.items():
            if block_type == 'text_blocks':
                for block in blocks:
                    doc.add_paragraph(block.text)
            elif block_type == 'title_blocks':
                for block in blocks:
                    doc.add_heading(block.text, level=1)
            elif block_type == 'list_blocks':
                for block in blocks:
                    doc.add_paragraph(block.text, style='ListBullet')
            elif block_type == 'table_blocks':
                for block in blocks:
                    table = doc.add_table(rows=1, cols=1)
                    cell = table.cell(0, 0)
                    cell.text = block.text
            # need to implement for figure
        doc.save('demo.docx')
        return doc
            
    